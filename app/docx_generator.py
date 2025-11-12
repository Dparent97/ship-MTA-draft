from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from app.models import WorkItem
from flask import current_app
import os


def generate_docx(work_item_id):
    """
    Generate a .docx file matching the template format.
    Returns the filepath of the generated document.
    """
    from app.models import WorkItem

    work_item = WorkItem.query.get_or_404(work_item_id)

    # Create document
    doc = Document()

    # Set default font
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Calibri'
    font.size = Pt(11)

    # Title
    title = doc.add_paragraph()
    title_run = title.add_run('WORK ITEM DRAFT TEMPLATE')
    title_run.bold = True
    title_run.font.size = Pt(14)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    doc.add_paragraph()  # Blank line

    # Item NO.
    p = doc.add_paragraph()
    p.add_run('Item NO.: ').bold = True
    p.add_run(work_item.item_number)

    # Location
    p = doc.add_paragraph()
    p.add_run('Location: ').bold = True
    p.add_run(work_item.location)

    doc.add_paragraph()  # Blank line

    # Description heading
    heading = doc.add_paragraph()
    heading_run = heading.add_run('Description:')
    heading_run.bold = True
    heading_run.font.size = Pt(12)

    # Description content
    doc.add_paragraph(work_item.description)

    doc.add_paragraph()  # Blank line

    # Detail heading
    heading = doc.add_paragraph()
    heading_run = heading.add_run('Detail:')
    heading_run.bold = True
    heading_run.font.size = Pt(12)

    # Detail content
    doc.add_paragraph(work_item.detail)

    # Operator Furnished Material (if provided)
    if work_item.references:
        doc.add_paragraph()  # Blank line
        heading = doc.add_paragraph()
        heading_run = heading.add_run('Operator Furnished Material (OFM):')
        heading_run.bold = True
        heading_run.font.size = Pt(12)
        doc.add_paragraph(work_item.references)

    # Photos section
    doc.add_paragraph()  # Blank line
    heading = doc.add_paragraph()
    heading_run = heading.add_run('PHOTOS')
    heading_run.bold = True
    heading_run.font.size = Pt(12)

    # Add each photo
    for idx, photo in enumerate(work_item.photos, 1):
        doc.add_paragraph()  # Blank line

        # Photo
        photo_path = os.path.join(current_app.config['UPLOAD_FOLDER'], photo.filename)
        if os.path.exists(photo_path):
            try:
                doc.add_picture(photo_path, width=Inches(4))
            except Exception as e:
                doc.add_paragraph(f'[Error loading photo: {photo.filename}]')

        # Caption
        caption_p = doc.add_paragraph()
        caption_p.add_run(f'Photo {idx} Caption: ').italic = True
        caption_p.add_run(photo.caption)

    # Metadata footer
    doc.add_paragraph()
    doc.add_paragraph()
    footer = doc.add_paragraph()
    footer_run = footer.add_run(
        f'Submitted by: {work_item.submitter_name} | '
        f'Date: {work_item.submitted_at.strftime("%Y-%m-%d %H:%M")}'
    )
    footer_run.font.size = Pt(9)
    footer_run.font.color.rgb = RGBColor(128, 128, 128)

    # Save document
    filename = f"{work_item.item_number}_{work_item.description[:30].replace(' ', '_')}.docx"
    # Use absolute path from Flask app root
    docs_folder = os.path.join(os.path.dirname(current_app.root_path), current_app.config['GENERATED_DOCS_FOLDER'])
    os.makedirs(docs_folder, exist_ok=True)
    filepath = os.path.join(docs_folder, filename)
    doc.save(filepath)
    
    return filepath


def generate_multiple_docx(work_item_ids):
    """
    Generate multiple .docx files and return list of filepaths.
    """
    filepaths = []
    for work_item_id in work_item_ids:
        try:
            filepath = generate_docx(work_item_id)
            filepaths.append(filepath)
        except Exception as e:
            print(f"Error generating document for work item {work_item_id}: {e}")
    return filepaths
