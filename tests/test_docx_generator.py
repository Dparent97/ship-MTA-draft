"""Tests for DOCX document generation."""
import pytest
import os
from docx import Document
from app.docx_generator import generate_docx, generate_multiple_docx
from app.models import WorkItem, Photo
from app import db


class TestGenerateDocx:
    """Test generate_docx function."""

    def test_generate_docx_success(self, app, init_database, sample_work_item):
        """Test successful document generation."""
        with app.app_context():
            filepath = generate_docx(sample_work_item.id)

            # Verify file was created
            assert os.path.exists(filepath)
            assert filepath.endswith('.docx')

            # Verify document content
            doc = Document(filepath)
            doc_text = '\n'.join([p.text for p in doc.paragraphs])

            assert 'WORK ITEM DRAFT TEMPLATE' in doc_text
            assert sample_work_item.item_number in doc_text
            assert sample_work_item.location in doc_text
            assert sample_work_item.description in doc_text
            assert sample_work_item.detail in doc_text

            # Cleanup
            if os.path.exists(filepath):
                os.remove(filepath)

    def test_generate_docx_with_photos(self, app, init_database, sample_work_item_with_photos):
        """Test document generation with photos."""
        with app.app_context():
            filepath = generate_docx(sample_work_item_with_photos.id)

            assert os.path.exists(filepath)

            # Verify document contains photo section
            doc = Document(filepath)
            doc_text = '\n'.join([p.text for p in doc.paragraphs])
            assert 'PHOTOS' in doc_text
            assert 'Photo 1 Caption' in doc_text

            # Cleanup
            if os.path.exists(filepath):
                os.remove(filepath)

    def test_generate_docx_with_references(self, app, init_database):
        """Test document generation with references/OFM."""
        with app.app_context():
            work_item = WorkItem(
                item_number='TEST_REF',
                location='Test Location',
                ns_equipment='Test Equipment',
                description='Test Description',
                detail='Test Detail',
                references='Part #12345, Manual page 67',
                submitter_name='DP'
            )
            db.session.add(work_item)
            db.session.commit()

            filepath = generate_docx(work_item.id)

            assert os.path.exists(filepath)

            # Verify document contains references
            doc = Document(filepath)
            doc_text = '\n'.join([p.text for p in doc.paragraphs])
            assert 'Operator Furnished Material (OFM)' in doc_text
            assert 'Part #12345' in doc_text

            # Cleanup
            if os.path.exists(filepath):
                os.remove(filepath)

    def test_generate_docx_without_references(self, app, init_database):
        """Test document generation without references."""
        with app.app_context():
            work_item = WorkItem(
                item_number='TEST_NO_REF',
                location='Test Location',
                ns_equipment='Test Equipment',
                description='Test Description',
                detail='Test Detail',
                references='',
                submitter_name='DP'
            )
            db.session.add(work_item)
            db.session.commit()

            filepath = generate_docx(work_item.id)

            assert os.path.exists(filepath)

            # Verify document does not contain OFM section
            doc = Document(filepath)
            doc_text = '\n'.join([p.text for p in doc.paragraphs])
            assert 'Operator Furnished Material (OFM)' not in doc_text

            # Cleanup
            if os.path.exists(filepath):
                os.remove(filepath)

    def test_generate_docx_includes_metadata(self, app, init_database, sample_work_item):
        """Test that document includes metadata footer."""
        with app.app_context():
            filepath = generate_docx(sample_work_item.id)

            assert os.path.exists(filepath)

            # Verify metadata
            doc = Document(filepath)
            doc_text = '\n'.join([p.text for p in doc.paragraphs])
            assert f'Submitted by: {sample_work_item.submitter_name}' in doc_text

            # Cleanup
            if os.path.exists(filepath):
                os.remove(filepath)

    def test_generate_docx_filename_format(self, app, init_database, sample_work_item):
        """Test that generated filename follows expected format."""
        with app.app_context():
            filepath = generate_docx(sample_work_item.id)

            filename = os.path.basename(filepath)
            # Should start with item number
            assert filename.startswith(sample_work_item.item_number)
            # Should end with .docx
            assert filename.endswith('.docx')

            # Cleanup
            if os.path.exists(filepath):
                os.remove(filepath)

    def test_generate_docx_creates_directory(self, app, init_database, sample_work_item):
        """Test that generate_docx creates output directory if needed."""
        with app.app_context():
            # Directory should be created by the function
            filepath = generate_docx(sample_work_item.id)

            # Verify directory exists
            directory = os.path.dirname(filepath)
            assert os.path.exists(directory)

            # Cleanup
            if os.path.exists(filepath):
                os.remove(filepath)

    def test_generate_docx_invalid_work_item_id(self, app, init_database):
        """Test error handling for invalid work item ID."""
        with app.app_context():
            with pytest.raises(Exception):  # Should raise 404 or similar
                generate_docx(99999)

    def test_generate_docx_with_missing_photo_file(self, app, init_database, sample_work_item):
        """Test handling of missing photo file."""
        with app.app_context():
            # Add photo record without actual file
            photo = Photo(
                filename='missing_photo.jpg',
                caption='Missing photo',
                work_item_id=sample_work_item.id
            )
            db.session.add(photo)
            db.session.commit()

            # Should still generate document
            filepath = generate_docx(sample_work_item.id)
            assert os.path.exists(filepath)

            # Cleanup
            if os.path.exists(filepath):
                os.remove(filepath)


class TestGenerateMultipleDocx:
    """Test generate_multiple_docx function."""

    def test_generate_multiple_docx_success(self, app, init_database):
        """Test generating multiple documents."""
        with app.app_context():
            # Create multiple work items
            work_items = []
            for i in range(3):
                work_item = WorkItem(
                    item_number=f'TEST_{i:03d}',
                    location=f'Location {i}',
                    ns_equipment=f'Equipment {i}',
                    description=f'Description {i}',
                    detail=f'Detail {i}',
                    submitter_name='DP'
                )
                db.session.add(work_item)
                work_items.append(work_item)
            db.session.commit()

            # Generate documents
            work_item_ids = [wi.id for wi in work_items]
            filepaths = generate_multiple_docx(work_item_ids)

            # Verify all documents were created
            assert len(filepaths) == 3
            for filepath in filepaths:
                assert os.path.exists(filepath)

            # Cleanup
            for filepath in filepaths:
                if os.path.exists(filepath):
                    os.remove(filepath)

    def test_generate_multiple_docx_empty_list(self, app, init_database):
        """Test generating documents with empty list."""
        with app.app_context():
            filepaths = generate_multiple_docx([])
            assert filepaths == []

    def test_generate_multiple_docx_with_errors(self, app, init_database, sample_work_item):
        """Test handling errors in batch generation."""
        with app.app_context():
            # Mix valid and invalid IDs
            work_item_ids = [sample_work_item.id, 99999, 99998]
            filepaths = generate_multiple_docx(work_item_ids)

            # Should succeed for valid ID, skip invalid ones
            assert len(filepaths) >= 1

            # Cleanup
            for filepath in filepaths:
                if os.path.exists(filepath):
                    os.remove(filepath)

    def test_generate_multiple_docx_preserves_order(self, app, init_database):
        """Test that documents are generated in order."""
        with app.app_context():
            # Create work items
            work_items = []
            for i in range(3):
                work_item = WorkItem(
                    item_number=f'ORDER_{i}',
                    location='Test',
                    ns_equipment='Test',
                    description='Test',
                    detail='Test',
                    submitter_name='DP'
                )
                db.session.add(work_item)
                work_items.append(work_item)
            db.session.commit()

            work_item_ids = [wi.id for wi in work_items]
            filepaths = generate_multiple_docx(work_item_ids)

            # Verify order is maintained
            for i, filepath in enumerate(filepaths):
                assert f'ORDER_{i}' in filepath

            # Cleanup
            for filepath in filepaths:
                if os.path.exists(filepath):
                    os.remove(filepath)
